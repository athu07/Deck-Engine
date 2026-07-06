# -*- coding: utf-8 -*-
"""
research.py  --  Read an uploaded "deep research" file (PDF or text) into plain
text so the matcher and the rationale/gap analysis can use it alongside the
meeting notes.

Kept deliberately simple and OFFLINE: it only extracts text that is already in
the file. A scanned / image-only PDF has no selectable text, so it returns ""
(the caller warns the user). Nothing is sent anywhere.
"""

import io


def extract_text(file_storage, max_chars=60000):
    """Text from a Werkzeug FileStorage (the uploaded research file), or ""."""
    if not file_storage or not getattr(file_storage, "filename", ""):
        return ""
    name = file_storage.filename.lower()
    try:
        data = file_storage.read()
    except Exception:
        return ""
    if not data:
        return ""
    if name.endswith(".pdf"):
        text = _pdf_text(data)
    elif name.endswith(".docx"):
        text = _docx_text(data)
    else:                                   # .txt / .md / anything text-ish
        text = data.decode("utf-8", errors="ignore")
    text = (text or "").strip()
    return text[:max_chars]


def _docx_text(data):
    """Extract text (paragraphs + table cells) from a .docx byte string."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _pdf_text(data):
    """Extract text from a PDF byte string. Tries pypdf, then PyMuPDF (fitz)."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        out = "\n".join((page.extract_text() or "") for page in reader.pages)
        if out.strip():
            return out
    except Exception:
        pass
    try:                                    # fallback: PyMuPDF is better on some PDFs
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    except Exception:
        return ""
